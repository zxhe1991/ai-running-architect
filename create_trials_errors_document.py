"""
创建项目开发过程中的关键试错与错误总结文档
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add heading with Chinese font support"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph_with_font(doc, text, bold=False):
    """Add paragraph with Chinese font support"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if bold:
            run.bold = True
    return para

def add_code_block(doc, code_text):
    """Add code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    para.style = 'List Paragraph'
    return para

def create_trials_errors_document():
    """Create comprehensive trials and errors document"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('AI Running Architect 项目开发过程中的关键试错与错误总结', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(24)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Subtitle
    subtitle = doc.add_paragraph(f'创建日期: {datetime.now().strftime("%Y年%m月%d日")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_paragraph()
    
    # Introduction
    add_paragraph_with_font(doc, 
        '本文档总结了 AI Running Architect 项目开发过程中遇到的关键问题、尝试的解决方案以及最终的处理方法。'
        '这些经验教训对于未来项目的开发和维护具有重要参考价值。')
    
    doc.add_page_break()
    
    # Issue 1: API Empty Response
    add_heading_with_style(doc, '1. API 返回空响应问题（最严重）', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_paragraph_with_font(doc, 
        '• API 调用成功，返回 500-2000 completion tokens')
    add_paragraph_with_font(doc, 
        '• 但 message.content 字段始终为空字符串')
    add_paragraph_with_font(doc, 
        '• 无论使用什么配置（json_object、流式、直接 HTTP）都相同')
    
    add_heading_with_style(doc, '尝试的解决方案：', 2)
    solutions_tried = [
        '使用 OpenAI SDK 标准调用',
        '使用 json_object 格式',
        '直接 HTTP 请求',
        '流式响应',
        '减少 max_tokens',
        '不使用 system prompt',
        '测试多个模型（gpt-4, gpt-3.5-turbo, claude-3 系列）— 均不支持'
    ]
    for solution in solutions_tried:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '最终解决方案：', 2)
    final_solutions = [
        '创建 mock_coach_response.py 模块提供模拟响应',
        '添加 USE_MOCK_RESPONSE 环境变量开关',
        '在 .env 文件中设置 USE_MOCK_RESPONSE=true',
        '结论：这是 API 端点的 bug，需要 space.ai-builders.com 修复'
    ]
    for solution in final_solutions:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons = [
        '外部 API 问题需要备选方案',
        '模拟响应有助于继续开发和测试',
        '需要定期监控 API 状态'
    ]
    for lesson in lessons:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Issue 2: StringDtype Compatibility
    add_heading_with_style(doc, '2. Pandas StringDtype 兼容性问题', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_code_block(doc, 'Error loading pickle file: (<StringDtype(storage=\'python\', na_value=nan)>, array([\'Running\', ...], dtype=object))')
    add_paragraph_with_font(doc, 
        '• 使用 Pandas 2.0+ 的 StringDtype 保存的 DataFrame')
    add_paragraph_with_font(doc, 
        '• 旧版本或不同环境加载 pickle 文件时失败')
    add_paragraph_with_font(doc, 
        '• 导致历史数据搜索功能完全无法使用')
    
    add_heading_with_style(doc, '尝试的解决方案：', 2)
    solutions_tried_2 = [
        '直接修复 pickle 文件（临时）',
        '从 CSV 重新加载并转换'
    ]
    for solution in solutions_tried_2:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '最终解决方案：', 2)
    add_paragraph_with_font(doc, 
        '在 build_index.py 中，保存前将所有字符串列转换为 object dtype：')
    add_code_block(doc, '''# Convert all string columns to object dtype before saving
for col in df.columns:
    if df[col].dtype.name == 'string' or 'StringDtype' in str(df[col].dtype):
        df[col] = df[col].astype('object')''')
    
    add_paragraph_with_font(doc, 
        '• 在 app.py 的 search_similar_runs 函数中添加错误处理和自动修复逻辑')
    add_paragraph_with_font(doc, 
        '• 创建 fix_pickle_file.py 工具脚本用于修复已损坏的文件')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons_2 = [
        '使用兼容性更好的数据类型（object 而非 StringDtype）',
        '添加错误处理和自动恢复机制',
        '提供修复工具脚本'
    ]
    for lesson in lessons_2:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Issue 3: Historical Data Display
    add_heading_with_style(doc, '3. 历史数据索引显示问题', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_paragraph_with_font(doc, 
        '• 用户报告"看不到历史数据索引的结果"')
    add_paragraph_with_font(doc, 
        '• 索引已构建，但界面不显示历史对比')
    add_paragraph_with_font(doc, 
        '• 搜索功能可能正常工作，但用户看不到反馈')
    
    add_heading_with_style(doc, '尝试的解决方案：', 2)
    solutions_tried_3 = [
        '添加调试信息和状态显示',
        '检查索引文件是否存在',
        '验证搜索查询构建逻辑'
    ]
    for solution in solutions_tried_3:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '最终解决方案：', 2)
    final_solutions_3 = [
        '添加更明确的状态消息（"✓ 找到 X 条相似的历史跑步记录"）',
        '改进 knowledge_base_built 检查逻辑，不仅检查文件存在，还验证索引有效性',
        '添加警告消息指导用户',
        '创建 WHY_HISTORICAL_DATA.md 和 HISTORICAL_DATA_EXPLANATION.md 文档'
    ]
    for solution in final_solutions_3:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons_3 = [
        '用户反馈很重要，需要清晰的视觉反馈',
        '状态检查要更严格（不仅检查文件存在）',
        '提供文档帮助用户理解功能'
    ]
    for lesson in lessons_3:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Issue 4: Vertical Oscillation
    add_heading_with_style(doc, '4. 垂直振荡（Vertical Oscillation）数据缺失', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_paragraph_with_font(doc, 
        '• 用户报告"为什么没有垂直振荡信息"')
    add_paragraph_with_font(doc, 
        '• CSV 分析器返回的数据键名不匹配')
    add_paragraph_with_font(doc, 
        '• csv_analyzer.py 返回 \'vertical_oscillation\'，但 app.py 期望 \'vertical_oscillation_metrics\'')
    
    add_heading_with_style(doc, '尝试的解决方案：', 2)
    solutions_tried_4 = [
        '检查 CSV 列名匹配逻辑',
        '改进垂直振荡列的数据清理'
    ]
    for solution in solutions_tried_4:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '最终解决方案：', 2)
    final_solutions_4 = [
        '统一键名为 \'vertical_oscillation_metrics\'',
        '改进列匹配逻辑，处理变体（如 "Avg Vertical Oscillationcm"）',
        '清理垂直振荡列，将 \'--\' 替换为空字符串并转换为数值'
    ]
    for solution in final_solutions_4:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons_4 = [
        '保持数据键名一致性',
        '处理数据格式变体',
        '改进数据清理逻辑'
    ]
    for lesson in lessons_4:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Issue 5: Excel Export
    add_heading_with_style(doc, '5. Excel 导出格式问题', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_paragraph_with_font(doc, 
        '• 用户报告"Excel 文件无效"')
    add_paragraph_with_font(doc, 
        '• 初始实现可能仍使用 PDF 生成逻辑')
    add_paragraph_with_font(doc, 
        '• Excel 文件无法正常打开')
    
    add_heading_with_style(doc, '尝试的解决方案：', 2)
    solutions_tried_5 = [
        '检查 Excel 生成代码',
        '验证 openpyxl 库是否正确使用'
    ]
    for solution in solutions_tried_5:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '最终解决方案：', 2)
    final_solutions_5 = [
        '重写 Excel 导出逻辑，使用 pandas.ExcelWriter',
        '创建多个工作表（Key Metrics, Detailed Analysis, Immediate Assessment, Training Plan, Training Strategy）',
        '将长文本内容按换行符分割成多行，提高可读性',
        '确保每个工作表格式正确'
    ]
    for solution in final_solutions_5:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons_5 = [
        '彻底重写有问题的功能，而不是修补',
        '测试导出文件的可用性',
        '考虑用户体验（可读性）'
    ]
    for lesson in lessons_5:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Issue 6: Deployment
    add_heading_with_style(doc, '6. 部署相关问题', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_paragraph_with_font(doc, 
        '• Dockerfile 中 PORT 环境变量扩展问题')
    add_paragraph_with_font(doc, 
        '• 部署配置中的服务名称和 URL 设置')
    add_paragraph_with_font(doc, 
        '• GitHub 仓库设置和代码推送')
    
    add_heading_with_style(doc, '尝试的解决方案：', 2)
    solutions_tried_6 = [
        '使用 shell 形式的 CMD 指令确保环境变量正确扩展',
        '验证部署配置格式'
    ]
    for solution in solutions_tried_6:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '最终解决方案：', 2)
    add_paragraph_with_font(doc, 
        'Dockerfile 使用：')
    add_code_block(doc, 'CMD sh -c "streamlit run app.py --server.port=${PORT:-8501}"')
    final_solutions_6 = [
        '创建 deploy-config.json 统一管理部署配置',
        '创建 DEPLOYMENT_GUIDE.md 详细说明部署步骤',
        '提供部署状态检查脚本'
    ]
    for solution in final_solutions_6:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons_6 = [
        'Docker 环境变量扩展需要 shell 形式',
        '提供详细的部署文档',
        '提供状态检查工具'
    ]
    for lesson in lessons_6:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Issue 7: Index Building Timing
    add_heading_with_style(doc, '7. 历史数据索引构建时机问题', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_paragraph_with_font(doc, 
        '• 用户希望只在点击"分析并获取建议"时才构建索引')
    add_paragraph_with_font(doc, 
        '• 而不是在上传 CSV 文件时自动构建')
    
    add_heading_with_style(doc, '解决方案：', 2)
    solutions_7 = [
        '将索引构建逻辑从文件上传处理移到"分析"按钮点击处理',
        '添加条件检查：只有当 CSV 存在且索引未构建时才构建'
    ]
    for solution in solutions_7:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons_7 = [
        '理解用户工作流程需求',
        '延迟执行可以改善用户体验'
    ]
    for lesson in lessons_7:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Issue 8: Syntax Error
    add_heading_with_style(doc, '8. 语法错误（SyntaxError）', 1)
    
    add_heading_with_style(doc, '问题描述：', 2)
    add_code_block(doc, 'SyntaxError: expected \'except\' or \'finally\' block')
    add_paragraph_with_font(doc, 
        '• pace 变量赋值位置错误')
    add_paragraph_with_font(doc, 
        '• 在 try 块外赋值导致语法错误')
    
    add_heading_with_style(doc, '解决方案：', 2)
    solutions_8 = [
        '将变量赋值移到正确的 try 块内',
        '确保 try-except 块结构正确'
    ]
    for solution in solutions_8:
        para = doc.add_paragraph(solution, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '经验教训：', 2)
    lessons_8 = [
        '仔细检查代码结构',
        '使用测试脚本发现语法错误'
    ]
    for lesson in lessons_8:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Summary
    add_heading_with_style(doc, '总结：关键经验教训', 1)
    
    summary_lessons = [
        '外部依赖问题需要备选方案（API 空响应 → 模拟响应）',
        '数据兼容性很重要（StringDtype → object dtype）',
        '用户反馈驱动改进（历史数据显示、Excel 格式）',
        '错误处理与自动恢复（pickle 文件修复）',
        '文档很重要（部署指南、问题解决文档）',
        '测试驱动开发（创建多个测试脚本）',
        '代码一致性（键名统一、命名规范）',
        '用户体验优先（清晰的反馈、合理的时机）'
    ]
    
    for lesson in summary_lessons:
        para = doc.add_paragraph(lesson, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_paragraph_with_font(doc, 
        '\n这些问题和解决方案已记录在项目文档中，便于未来参考和维护。')
    
    # Save document
    filename = f'AI_Running_Architect_试错总结_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f"✅ 试错总结文档已保存为: {filename}")
    return filename

if __name__ == "__main__":
    create_trials_errors_document()
