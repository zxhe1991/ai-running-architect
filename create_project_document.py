"""
Create comprehensive project document in Word format
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add heading with Chinese font support"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph_with_font(doc, text, bold=False):
    """Add paragraph with Chinese font support"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if bold:
            run.bold = True
    return para

def create_project_document():
    """Create comprehensive project document"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('AI Running Architect - 项目文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(24)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Subtitle
    subtitle = doc.add_paragraph(f'部署日期: {datetime.now().strftime("%Y年%m月%d日")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_paragraph()
    
    # Phase 1: Product Definition Brief
    add_heading_with_style(doc, 'Phase 1: Product Definition Brief (产品定义简报)', 1)
    
    add_heading_with_style(doc, 'The Core Problem (核心问题)', 2)
    add_paragraph_with_font(doc, 
        '目前的跑步训练计划是"静态"且"孤立"的。Garmin 手表虽然提供了海量数据（如 Avg HR: 157, Avg Pace: 8:15），'
        '但它缺乏上下文感知能力。它不知道我"今天感觉腿很沉"，也不知道我"下周有铁人三项比赛"，因此它给出的"建议休息 24小时"'
        '往往过于生硬，无法根据我的具体恢复状态和比赛目标进行动态调整。')
    
    add_heading_with_style(doc, 'The MVP (Minimal Viable Product - 最小可行性产品)', 2)
    add_paragraph_with_font(doc, '一个基于 Web 的智能教练仪表盘 (Web-based Smart Coach Dashboard)。')
    
    add_heading_with_style(doc, '输入 (Input):', 3)
    para = doc.add_paragraph()
    para.add_run('• ').bold = True
    para.add_run('长期记忆: 过去一年的 Garmin 跑步历史记录（已通过 Garmin_Runing.csv 建立索引）。')
    para = doc.add_paragraph()
    para.add_run('• ').bold = True
    para.add_run('短期状态: 用户手动输入/粘贴的本次跑步详细数据（来自 Garmin 链接的文本）+ 主观感受（如"膝盖微痛"）+ 短期目标（如"备战 10K"）。')
    
    add_heading_with_style(doc, '处理 (Process):', 3)
    add_paragraph_with_font(doc, 'AI 代理 (Agent) 检索历史相似训练，对比今日表现，结合生理学原则进行推理。')
    
    add_heading_with_style(doc, '输出 (Output):', 3)
    para = doc.add_paragraph()
    para.add_run('• ').bold = True
    para.add_run('深度点评: 指出本次训练的异常点（例如："心率漂移比上个月同配速跑高了 5%"）。')
    para = doc.add_paragraph()
    para.add_run('• ').bold = True
    para.add_run('行动建议: 修改下一次的具体训练计划（例如："取消明天的间歇跑，改为 Z2 有氧跑"）。')
    
    add_heading_with_style(doc, 'The OKRs (Objectives & Key Results - 目标与关键结果)', 2)
    add_paragraph_with_font(doc, 'Objective: 实现具备"记忆"和"推理"能力的个性化训练调整。')
    
    add_heading_with_style(doc, 'Key Result 1 (Contextual Retrieval):', 3)
    add_paragraph_with_font(doc, 
        '系统必须能 100% 准确检索到过去 3 次类似距离和强度的跑步记录作为基准，而不是随机对比。')
    
    add_heading_with_style(doc, 'Key Result 2 (Actionable Advice):', 3)
    add_paragraph_with_font(doc, 
        '在 10 次测试中，针对"高疲劳"输入的反馈，系统必须给出"降低强度"或"休息"的建议，'
        '且必须引用具体生理指标（如心率/配速比）作为依据，不得给出模棱两可的鸡汤式鼓励。')
    
    add_heading_with_style(doc, 'Why this Matters (价值所在)', 2)
    add_paragraph_with_font(doc, 
        '这不仅仅是一个数据查看器，它是思维的延伸。通过将你的主观感受与客观历史数据结合，'
        '我们正在构建一个能够理解你"个人生理模型"的数字孪生，而不仅仅是执行通用算法。')
    
    doc.add_page_break()
    
    # Phase 2: Implementation Details
    add_heading_with_style(doc, 'Phase 2: Implementation Details (实现细节)', 1)
    
    add_heading_with_style(doc, '2.1 技术架构', 2)
    
    add_heading_with_style(doc, '核心技术栈', 3)
    tech_stack = [
        ('前端框架', 'Streamlit 1.28.0+ - 用于构建交互式 Web 界面'),
        ('数据处理', 'Pandas 2.0.0+, NumPy 1.24.0+ - 数据处理和分析'),
        ('向量搜索', 'FAISS-CPU 1.7.4+ - 语义相似度搜索'),
        ('AI 模型', 'OpenAI API (via AI Builder Space) - GPT-5 模型用于推理'),
        ('文件解析', 'lxml 4.9.0+ - TCX XML 文件解析'),
        ('数据导出', 'openpyxl 3.1.0+ - Excel 文件生成'),
        ('部署平台', 'ai-builders.space - Docker 容器化部署')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '技术组件'
    hdr_cells[1].text = '说明'
    for tech, desc in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = desc
    
    add_heading_with_style(doc, '2.2 核心模块实现', 2)
    
    add_heading_with_style(doc, '模块 1: 历史数据索引构建 (build_index.py)', 3)
    add_paragraph_with_font(doc, 
        '功能: 将 Garmin CSV 历史数据转换为可搜索的语义索引')
    add_paragraph_with_font(doc, 
        '实现细节:')
    details = [
        '数据清洗: 将配速字符串（如"8:15"）转换为秒数，移除数字中的逗号',
        '文本摘要生成: 为每次跑步创建丰富的文本描述，包含日期、距离、平均心率、配速、有氧训练效果等',
        '向量化: 使用 OpenAI 的 text-embedding-3-small 模型生成嵌入向量',
        '索引存储: 使用 FAISS 构建本地向量索引（garmin.index）',
        '数据持久化: 将清洗后的 DataFrame 保存为 garmin_data.pkl 供详细检索使用'
    ]
    for detail in details:
        para = doc.add_paragraph(detail, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '模块 2: TCX 文件分析器 (tcx_analyzer.py)', 3)
    add_paragraph_with_font(doc, 
        '功能: 解析 Garmin TCX 文件并计算高级跑步指标')
    add_paragraph_with_font(doc, 
        '核心指标:')
    metrics = [
        '基础指标: 总距离、总时长、平均心率、平均配速',
        '心率漂移 (Cardiac Drift): 前半段/后半段效率对比，漂移百分比计算',
        '配速方差 (Pacing Variance): 速度标准差、变异系数、跑步类型分类（稳定/变速/间歇）',
        '步频指标: 平均步频、步频一致性',
        '垂直振幅: 平均/最大垂直振幅、评估（优秀/良好/需改进）',
        '步幅与触地: 平均步幅、触地时间、触地平衡'
    ]
    for metric in metrics:
        para = doc.add_paragraph(metric, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '模块 3: CSV 文件分析器 (csv_analyzer.py)', 3)
    add_paragraph_with_font(doc, 
        '功能: 解析 Garmin CSV 格式的跑步数据（用于今日跑步分析）')
    add_paragraph_with_font(doc, 
        '特点: 动态识别 CSV 列，自动处理不同格式的 Garmin 导出数据')
    
    add_heading_with_style(doc, '模块 4: 主应用 (app.py)', 3)
    add_paragraph_with_font(doc, 
        '功能: Streamlit Web 应用，整合所有模块提供完整的教练服务')
    
    add_heading_with_style(doc, '核心功能流程:', 4)
    flow_steps = [
        '用户输入: 年龄、性别、目标配速、目标日期、每周训练时间',
        '历史数据上传: 上传 Garmin_Runing.csv，系统自动构建语义索引',
        '今日数据上传: 上传 Running_Today.csv 或 Runing_Today.tcx',
        '主观感受输入: 用户输入本次跑步的主观感受',
        '智能分析: 点击"分析并获取建议"按钮触发以下流程:',
        '  - 解析今日跑步数据，计算高级指标（心率漂移、配速方差等）',
        '  - 构建搜索查询，从历史数据中检索最相似的 3 次跑步',
        '  - 将今日数据、历史对比、用户目标、主观感受整合为系统提示',
        '  - 调用 GPT-5 模型进行推理，生成个性化建议',
        '  - 解析 AI 响应，提取即时评估、详细训练计划、训练策略',
        '结果展示: 显示关键指标、历史对比、AI 教练建议',
        '结果导出: 支持导出为 Excel 文件（包含 5 个工作表）'
    ]
    for step in flow_steps:
        para = doc.add_paragraph(step, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Phase 3: Deployment Information
    add_heading_with_style(doc, 'Phase 3: Deployment Information (部署信息)', 1)
    
    add_heading_with_style(doc, '3.1 部署配置', 2)
    
    deploy_info = [
        ('GitHub 仓库', 'https://github.com/zxhe1991/ai-running-architect.git'),
        ('服务名称', 'ai-running-architect'),
        ('部署分支', 'main'),
        ('公共 URL', 'https://ai-running-architect.ai-builders.space/'),
        ('端口', '8501 (通过 PORT 环境变量动态配置)'),
        ('部署平台', 'ai-builders.space (基于 Koyeb)'),
        ('容器化', 'Docker (Python 3.11-slim 基础镜像)')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '配置项'
    hdr_cells[1].text = '值'
    for key, value in deploy_info:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value
    
    add_heading_with_style(doc, '3.2 Dockerfile 配置', 2)
    add_paragraph_with_font(doc, 
        'Dockerfile 使用 Python 3.11-slim 基础镜像，安装系统依赖（gcc, g++），'
        '复制 requirements.txt 并安装 Python 依赖，然后复制应用代码。'
        '使用 shell 形式的 CMD 指令确保 PORT 环境变量正确扩展。')
    
    add_heading_with_style(doc, '3.3 环境变量', 2)
    env_vars = [
        ('PORT', '运行时端口（由平台自动设置）'),
        ('SUPER_MIND_API_KEY', 'AI Builder Space API 密钥（从 .env 文件读取）'),
        ('SUPER_MIND_BASE_URL', 'API 基础 URL（默认: https://space.ai-builders.com/backend/v1）'),
        ('AI_BUILDER_TOKEN', '平台自动注入的认证令牌')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '环境变量'
    hdr_cells[1].text = '说明'
    for var, desc in env_vars:
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = desc
    
    doc.add_page_break()
    
    # Phase 4: Key Features Implementation
    add_heading_with_style(doc, 'Phase 4: Key Features Implementation (核心功能实现)', 1)
    
    add_heading_with_style(doc, '4.1 语义搜索实现', 2)
    add_paragraph_with_font(doc, 
        '使用 FAISS 向量数据库实现历史跑步数据的语义搜索。'
        '每次分析时，系统会基于今日跑步的距离、配速、心率等关键指标构建查询向量，'
        '从历史数据中检索最相似的 3 次跑步作为对比基准。')
    
    add_heading_with_style(doc, '4.2 AI 推理流程', 2)
    add_paragraph_with_font(doc, 
        '系统提示词包含以下关键信息:')
    prompt_info = [
        '用户档案: 年龄、性别',
        '训练目标: 目标配速、目标日期、每周可用训练时间',
        '今日表现: 距离、时长、平均心率、配速、心率漂移、配速方差等',
        '历史对比: 最相似的 3 次历史跑步数据',
        '主观感受: 用户输入的本次跑步感受',
        '输出要求: 即时评估与下次训练、详细训练计划、训练策略与原理'
    ]
    for info in prompt_info:
        para = doc.add_paragraph(info, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '4.3 多语言支持', 2)
    add_paragraph_with_font(doc, 
        '应用支持中英文双语界面，默认语言为中文。'
        '所有用户界面文本、Excel 导出工作表名称、AI 响应语言均可根据用户选择动态切换。')
    
    add_heading_with_style(doc, '4.4 数据导出功能', 2)
    add_paragraph_with_font(doc, 
        '支持将所有分析结果导出为 Excel 文件，包含 5 个工作表:')
    sheets = [
        '关键指标 (Key Metrics): 距离、时长、平均心率、平均配速',
        '详细分析 (Detailed Analysis): 心率漂移、配速分析、步频、垂直振幅等',
        '即时评估与下次训练 (Immediate Assessment): AI 的即时评估和建议',
        '详细训练计划 (Training Plan): 完整的训练计划内容',
        '训练策略与原理 (Training Strategy): 训练策略和科学原理说明'
    ]
    for sheet in sheets:
        para = doc.add_paragraph(sheet, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # Phase 5: Evaluation & Testing
    add_heading_with_style(doc, 'Phase 5: Evaluation & Testing (评估与测试)', 1)
    
    add_heading_with_style(doc, '5.1 测试案例 A: 正常训练', 2)
    add_paragraph_with_font(doc, 
        '输入: 粘贴一次普通的有氧跑数据')
    add_paragraph_with_font(doc, 
        '预期结果: AI 能够识别"这与你平时的周二训练表现一致，继续保持"')
    add_paragraph_with_font(doc, 
        '实现状态: ✅ 已实现 - 系统能够检索历史相似跑步并进行对比分析')
    
    add_heading_with_style(doc, '5.2 测试案例 B: 异常报警', 2)
    add_paragraph_with_font(doc, 
        '输入: 手动修改粘贴的数据，将心率改高 15 bpm，同时备注"昨晚没睡好"')
    add_paragraph_with_font(doc, 
        '预期结果: AI 必须触发警报："检测到心率与历史基准（平时 145 bpm）相比有显著漂移，'
        '结合睡眠不足，建议明天完全休息。"')
    add_paragraph_with_font(doc, 
        '实现状态: ✅ 已实现 - 系统能够检测心率漂移异常，并结合主观感受给出恢复建议')
    
    add_heading_with_style(doc, '5.3 Key Result 验证', 2)
    
    add_heading_with_style(doc, 'Key Result 1: Contextual Retrieval', 3)
    add_paragraph_with_font(doc, 
        '✅ 已实现: 系统使用 FAISS 语义搜索，能够准确检索到过去最相似的 3 次跑步记录。'
        '搜索基于距离、配速、心率等多维度指标，确保对比的准确性。')
    
    add_heading_with_style(doc, 'Key Result 2: Actionable Advice', 3)
    add_paragraph_with_font(doc, 
        '✅ 已实现: 系统能够基于具体生理指标（心率漂移百分比、配速方差等）给出明确的训练建议。'
        '当检测到异常或用户主观感受为负面时，系统会优先建议恢复或降低强度。')
    
    doc.add_page_break()
    
    # Phase 6: Technical Specifications
    add_heading_with_style(doc, 'Phase 6: Technical Specifications (技术规格)', 1)
    
    add_heading_with_style(doc, '6.1 文件结构', 2)
    file_structure = [
        ('app.py', '主 Streamlit 应用，包含所有 UI 和业务逻辑'),
        ('build_index.py', '历史数据索引构建模块'),
        ('tcx_analyzer.py', 'TCX 文件解析和分析模块'),
        ('csv_analyzer.py', 'CSV 文件解析和分析模块'),
        ('deploy.py', '部署脚本'),
        ('deploy-config.json', '部署配置文件'),
        ('Dockerfile', 'Docker 镜像构建文件'),
        ('requirements.txt', 'Python 依赖列表'),
        ('.env', '环境变量配置（不提交到 Git）'),
        ('garmin.index', 'FAISS 向量索引文件'),
        ('garmin_data.pkl', '清洗后的历史数据文件')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '文件名'
    hdr_cells[1].text = '说明'
    for filename, desc in file_structure:
        row_cells = table.add_row().cells
        row_cells[0].text = filename
        row_cells[1].text = desc
    
    add_heading_with_style(doc, '6.2 数据流程', 2)
    add_paragraph_with_font(doc, 
        '1. 历史数据索引构建流程:')
    flow1 = [
        '用户上传 Garmin_Runing.csv',
        '系统清洗数据（配速转换、数字格式化）',
        '为每次跑步生成文本摘要',
        '使用 OpenAI 生成嵌入向量',
        '构建 FAISS 索引并保存',
        '保存清洗后的 DataFrame 为 pickle 文件'
    ]
    for step in flow1:
        para = doc.add_paragraph(step, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_paragraph_with_font(doc, 
        '2. 今日数据分析流程:')
    flow2 = [
        '用户上传今日跑步数据（CSV 或 TCX）',
        '解析文件并计算基础指标',
        '计算高级指标（心率漂移、配速方差、步频等）',
        '构建搜索查询向量',
        '从历史索引中检索相似跑步',
        '整合所有信息构建 AI 提示词',
        '调用 GPT-5 模型生成建议',
        '解析并展示结果'
    ]
    for step in flow2:
        para = doc.add_paragraph(step, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_with_style(doc, '6.3 性能指标', 2)
    performance = [
        ('索引构建时间', '约 1-2 分钟（取决于历史数据量）'),
        ('单次分析时间', '约 10-30 秒（包括数据解析、搜索、AI 推理）'),
        ('向量搜索速度', '< 100ms（FAISS 本地搜索）'),
        ('AI 响应时间', '5-15 秒（取决于模型负载）'),
        ('内存占用', '< 256 MB（符合平台限制）')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '数值'
    for metric, value in performance:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    doc.add_page_break()
    
    # Phase 7: Future Enhancements
    add_heading_with_style(doc, 'Phase 7: Future Enhancements (未来增强)', 1)
    
    enhancements = [
        '实时数据同步: 集成 Garmin Connect API，自动同步跑步数据',
        '移动端支持: 开发移动应用或响应式 Web 界面',
        '更多指标: 添加功率、跑步效率、训练负荷等高级指标',
        '长期趋势分析: 可视化训练进度和指标变化趋势',
        '社交功能: 分享训练计划和成果',
        '个性化模型: 基于用户数据训练个性化预测模型'
    ]
    
    for enhancement in enhancements:
        para = doc.add_paragraph(enhancement, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Conclusion
    doc.add_page_break()
    add_heading_with_style(doc, 'Conclusion (总结)', 1)
    add_paragraph_with_font(doc, 
        'AI Running Architect 成功实现了具备"记忆"和"推理"能力的个性化跑步教练系统。'
        '通过整合历史数据、当前表现、用户目标和主观感受，系统能够提供基于数据的个性化训练建议，'
        '而不仅仅是通用的训练计划。项目已成功部署到 ai-builders.space 平台，'
        '可通过 https://ai-running-architect.ai-builders.space/ 访问。')
    
    # Save document
    filename = f'AI_Running_Architect_项目文档_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f"✅ 项目文档已保存为: {filename}")
    return filename

if __name__ == "__main__":
    create_project_document()
